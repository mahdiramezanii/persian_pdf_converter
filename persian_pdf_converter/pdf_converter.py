import os

from tqdm import tqdm
from pdf2image import convert_from_path
import pytesseract
from PIL import Image
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import tempfile
from persian_pdf_converter import _my_random_string

from pathlib import Path

parent_path= Path(__file__).resolve().parent.parent

def pdf_to_word(pdf_path: str, output_dir: str, lang="fas+eng", **kwargs):
    output_dir = output_dir.replace("\\", "/")
    pdf_name = f"word-{_my_random_string(6)}"


    pages = convert_from_path(pdf_path,poppler_path=f"{parent_path}/statics/poppler-24.02/bin")

    pytesseract.pytesseract.tesseract_cmd = f"{parent_path}/statics/Tesseract-OCR/tesseract.exe"
    texts = []

    for i, page in tqdm(enumerate(pages), position=0):
        with tempfile.TemporaryDirectory() as img_dir:
            img_name = f'{pdf_name}-{i+1}.jpg'
            img_path = Path(img_dir) / img_name


            page.save(img_path, 'JPEG')
            text = pytesseract.image_to_string(Image.open(img_path), lang=lang)
            texts.append(text)

    document = Document()
    style_normal = document.styles['Normal']
    font = style_normal.font
    font.name = 'Arial'
    font.rtl = True

    style_h1 = document.styles['Heading 1']
    font = style_h1.font
    font.name = 'Arial'
    font.rtl = True

    for i, text in tqdm(enumerate(texts), position=0):
        heading = document.add_heading(f'صفحه: {i+1}', level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        heading.style = document.styles['Heading 1']

        paragraph = document.add_paragraph(text)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph.style = document.styles['Normal']

    output_path = Path(output_dir) / f'{pdf_name}.docx'
    document.save(output_path)

    return f'{pdf_name}.docx'


